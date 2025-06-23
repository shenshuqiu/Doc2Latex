"""
文档处理器

负责读取和处理 Word 文档，构建文档树结构。
"""

import os
import docx
from collections import OrderedDict
from treelib import Tree, Node
from typing import Dict, List, Optional

from ..utils.file_utils import (
    save_docx_to_dict, get_file_list, sort_by_serial, clean_directory
)
from ..utils.text_utils import check_chinese_square_brackets_pairs
from ..utils.doc_converter import convert_doc_to_docx
from ..config.settings import PATHS


class DocumentProcessor:
    """文档处理器类"""
    
    def __init__(self, raw_document_dir: str = None, processed_document_dir: str = None):
        """
        初始化文档处理器
        
        Args:
            raw_document_dir: 原始文档目录
            processed_document_dir: 处理后文档目录
        """
        self.raw_document_dir = raw_document_dir or str(PATHS["raw_document"])
        self.processed_document_dir = processed_document_dir or str(PATHS["document"])
        self.document_dict = OrderedDict()
        self.book_tree = Tree()
    
    def convert_all_doc_to_docx(self) -> None:
        """
        将所有 doc 文件转换为 docx 文件
        """
        file_list = get_file_list(self.raw_document_dir, ['.doc', '.docx'])
        
        for file in file_list:
            if file.endswith('.doc'):
                doc_file_path = os.path.join(self.raw_document_dir, file)
                docx_file_path = os.path.join(self.raw_document_dir, file.replace('.doc', '.docx'))
                
                try:
                    convert_doc_to_docx(doc_file_path, docx_file_path)
                    print(f"{file} 已被转换为 docx 文件")
                except Exception as e:
                    print(f"转换 {file} 时出错: {e}")
    
    def load_all_documents(self) -> None:
        """
        加载所有文档到字典中
        """
        file_list = get_file_list(self.raw_document_dir, ['.docx'])
        file_list = sort_by_serial(file_list)
        
        for file in file_list:
            docx_file_path = os.path.join(self.raw_document_dir, file)
            save_docx_to_dict(docx_file_path, self.document_dict)
    
    def validate_document_brackets(self) -> None:
        """
        验证文档中的中文方括号配对
        """
        for serial, content in self.document_dict.items():
            if content.get("text"):
                for paragraph_num, paragraph in enumerate(content["text"]):
                    if not check_chinese_square_brackets_pairs(paragraph):
                        raise AssertionError(f"方括号未配对: {serial} 第{paragraph_num + 1}段")
    
    def build_document_tree(self) -> None:
        """
        构建文档树结构并重新编号
        """
        # 创建根节点
        self.book_tree.create_node(identifier="root")
        
        # 章节计数器
        chapter_counter = 0
        section_counter = 0
        subsection_counter = 0
        
        for serial in self.document_dict.keys():
            doc_data = self.document_dict[serial]
            
            if doc_data["subsection"] == 0:
                if doc_data["section"] == 0:
                    if doc_data["chapter"] == 0:
                        raise NameError(f"文档命名错误：{serial}")
                    else:
                        # 新章节
                        chapter_counter += 1
                        section_counter = 0
                        new_tag = f"{chapter_counter}-0-0"
                        
                        self.book_tree.create_node(
                            identifier=serial,
                            data=doc_data,
                            tag=new_tag,
                            parent="root"
                        )
                else:
                    # 新节
                    section_counter += 1
                    subsection_counter = 0
                    new_tag = f"{chapter_counter}-{section_counter}-0"
                    
                    # 验证父文档存在
                    father_serial = f"{doc_data['chapter']}-0-0"
                    if father_serial not in self.document_dict:
                        raise AssertionError(f"{serial} 文档缺少父文档 {father_serial}")
                    
                    self.book_tree.create_node(
                        identifier=serial,
                        data=doc_data,
                        tag=new_tag,
                        parent=father_serial
                    )
            else:
                # 新小节
                subsection_counter += 1
                new_tag = f"{chapter_counter}-{section_counter}-{subsection_counter}"
                
                # 验证父文档存在
                father_serial = f"{doc_data['chapter']}-{doc_data['section']}-0"
                if father_serial not in self.document_dict:
                    raise AssertionError(f"{serial} 文档缺少父文档 {father_serial}")
                
                self.book_tree.create_node(
                    identifier=serial,
                    data=doc_data,
                    tag=new_tag,
                    parent=father_serial
                )
    
    def save_processed_documents(self) -> None:
        """
        保存处理后的文档到指定目录
        """
        # 清空目标目录
        clean_directory(self.processed_document_dir)
        
        # 遍历树结构保存文档
        for node in self.book_tree.expand_tree(mode=Tree.DEPTH, sorting=False):
            if node == self.book_tree.root:
                continue
            
            # 获取新的序号和数据
            new_serial = self.book_tree[node].tag
            doc_data = self.book_tree[node].data
            
            # 创建新文档
            new_document = docx.Document()
            
            # 设置字体格式
            style = new_document.styles['Normal']
            style.font.name = '宋体'
            style._element.rPr.rFonts.set(
                docx.oxml.shared.qn('w:eastAsia'), '宋体'
            )
            style.font.size = docx.shared.Pt(10.5)
            style.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
            
            # 添加标题
            new_document.add_paragraph(doc_data["name"])
            
            # 添加正文内容
            if doc_data.get("text"):
                for paragraph in doc_data["text"]:
                    new_document.add_paragraph(paragraph)
            
            # 保存文档
            new_docx_path = os.path.join(self.processed_document_dir, f"{new_serial}.docx")
            new_document.save(new_docx_path)
    
    def process_all_documents(self) -> None:
        """
        处理所有文档的完整流程
        """
        print("开始转换 doc 文件...")
        self.convert_all_doc_to_docx()
        
        print("加载所有文档...")
        self.load_all_documents()
        
        print("验证文档格式...")
        try:
            self.validate_document_brackets()
            print("文档格式验证通过")
        except AssertionError as e:
            print(f"文档格式验证失败: {e}")
            raise
        
        print("构建文档树...")
        self.build_document_tree()
        
        print("保存处理后的文档...")
        self.save_processed_documents()
        
        print("文档处理完成!")
    
    def get_document_dict(self) -> Dict:
        """
        获取文档字典
        
        Returns:
            文档字典
        """
        return self.document_dict
    
    def get_document_tree(self) -> Tree:
        """
        获取文档树
        
        Returns:
            文档树对象
        """
        return self.book_tree